import re
import sys
import io
from pathlib import Path
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
import multiprocessing

import spacy
import easyocr
import cv2
import numpy as np

from PIL import Image
from tqdm import tqdm
from docx import Document


class PentestAnonymiser:

    def __init__(self, custom_terms=None):
        # NLP model for entity recognition
        self.nlp = spacy.load("en_core_web_sm")

        # user-supplied masking terms (custom organisation names)
        self.detected_terms = set()
        if custom_terms:
            for t in custom_terms.split("|"):
                self.detected_terms.add(t.strip())

        # technology whitelist to avoid masking
        self.whitelist = {
            "javascript","java","python","php","node","react",
            "sql","mysql","postgres","mongodb",
            "linux","windows","apache","nginx",
            "http","https","tls","ssh",
            "owasp","xss","csrf","sqli",
            "cross-site","scripting","example","test"
        }

        # regex patterns
        self.ip_pattern = re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b')
        self.email_pattern = re.compile(r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\b')
        self.url_pattern = re.compile(r'https?://[^\s]+')
        self.domain_pattern = re.compile(r'\b(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}\b')
        self.port_pattern = re.compile(r'\bport\s+\d{1,5}\b', re.IGNORECASE)

        # hostname detection patterns
        self.hostname_patterns = [
            re.compile(r'\b[A-Z]{2,}-[A-Z0-9\-]{2,}\b'),
            re.compile(r'\b[a-z]+-(node|server|worker|runner|build)-\d+\b'),
            re.compile(r'\b[a-z]+-(prod|dev|test|stage)-\d+\b')
        ]

        # EasyOCR reader for Python-only OCR
        self.ocr_reader = easyocr.Reader(['en'], gpu=False)

    # -------------------------
    # Private IP subnet masking
    # -------------------------
    def mask_ip(self, ip):
        parts = ip.split(".")
        first = int(parts[0])
        second = int(parts[1])
        if first == 10:
            return "10.x.x.x"
        if first == 172 and 16 <= second <= 31:
            return "172.x.x.x"
        if first == 192 and second == 168:
            return "192.168.x.x"
        return "x.x.x.x"

    # -------------------------
    # Extract all text blocks
    # -------------------------
    def extract_text(self, doc):
        blocks = []
        for p in doc.paragraphs:
            blocks.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        blocks.append(p.text)
        for section in doc.sections:
            for p in section.header.paragraphs:
                blocks.append(p.text)
            for p in section.footer.paragraphs:
                blocks.append(p.text)
        return blocks

    # -------------------------
    # Domain & org detection
    # -------------------------
    def detect_domains(self, blocks):
        for text in blocks:
            for d in self.domain_pattern.findall(text):
                self.detected_terms.add(d)
                root = d.split(".")[0]
                if len(root) > 3:
                    self.detected_terms.add(root)

    # -------------------------
    # Hostname detection
    # -------------------------
    def detect_hostnames(self, blocks):
        for text in blocks:
            for pattern in self.hostname_patterns:
                matches = pattern.findall(text)
                for m in matches:
                    if isinstance(m, tuple):
                        m = m[0]
                    if m.lower() not in self.whitelist:
                        self.detected_terms.add(m)

    # -------------------------
    # Organisation detection via NLP
    # -------------------------
    def detect_orgs(self, blocks):
        for text in blocks:
            doc = self.nlp(text)
            for ent in doc.ents:
                if ent.label_ != "ORG":
                    continue
                name = ent.text.strip()
                if name.lower() in self.whitelist:
                    continue
                if len(name) < 3:
                    continue
                self.detected_terms.add(name)

    # -------------------------
    # Frequently repeated capitalised terms
    # -------------------------
    def detect_frequent_terms(self, blocks):
        words = []
        for text in blocks:
            words += re.findall(r'\b[A-Z][A-Za-z0-9]+\b', text)
        freq = Counter(words)
        for w, count in freq.items():
            if count >= 3 and w.lower() not in self.whitelist:
                self.detected_terms.add(w)

    # -------------------------
    # Generic masking function
    # -------------------------
    def mask(self, text):
        return "x" * len(text)

    def anonymise_text(self, text):
        text = self.ip_pattern.sub(lambda m: self.mask_ip(m.group()), text)
        text = self.url_pattern.sub(lambda m: self.mask(m.group()), text)
        text = self.email_pattern.sub(lambda m: self.mask(m.group()), text)
        text = self.domain_pattern.sub(lambda m: self.mask(m.group()), text)
        text = self.port_pattern.sub(lambda m: self.mask(m.group()), text)
        for term in sorted(self.detected_terms, key=len, reverse=True):
            pattern = re.compile(re.escape(term), re.IGNORECASE)
            text = pattern.sub(lambda m: self.mask(m.group()), text)
        return text

    # -------------------------
    # Process paragraphs, tables, headers/footers
    # -------------------------
    def process_paragraphs(self, paragraphs):
        for p in paragraphs:
            for run in p.runs:
                if run.text:
                    run.text = self.anonymise_text(run.text)

    def process_tables(self, tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    self.process_paragraphs(cell.paragraphs)
                    self.process_tables(cell.tables)

    def process_headers(self, doc):
        for section in doc.sections:
            self.process_paragraphs(section.header.paragraphs)
            self.process_tables(section.header.tables)
            self.process_paragraphs(section.footer.paragraphs)
            self.process_tables(section.footer.tables)

    # -------------------------
    # Image sanitisation using EasyOCR
    # -------------------------
    def sanitise_image(self, image_bytes):
        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        img = np.array(image)
        results = self.ocr_reader.readtext(img)
        for bbox, text, conf in results:
            text = text.strip()
            if not text:
                continue
            if (
                self.ip_pattern.match(text) or
                self.domain_pattern.match(text) or
                self.email_pattern.match(text)
            ):
                x_min = min([pt[0] for pt in bbox])
                y_min = min([pt[1] for pt in bbox])
                x_max = max([pt[0] for pt in bbox])
                y_max = max([pt[1] for pt in bbox])
                cv2.rectangle(img, (int(x_min), int(y_min)), (int(x_max), int(y_max)), (0, 0, 0), -1)
        redacted = Image.fromarray(img)
        output = io.BytesIO()
        redacted.save(output, format="PNG")
        return output.getvalue()

    def process_images(self, doc):
        for rel in doc.part._rels:
            rel = doc.part._rels[rel]
            if "image" in rel.target_ref:
                image_part = rel.target_part
                try:
                    sanitised = self.sanitise_image(image_part.blob)
                    image_part._blob = sanitised
                except Exception:
                    pass

    # -------------------------
    # Anonymise single file
    # -------------------------
    def anonymise_file(self, path):
        doc = Document(path)
        blocks = self.extract_text(doc)
        self.detect_domains(blocks)
        self.detect_hostnames(blocks)
        self.detect_orgs(blocks)
        self.detect_frequent_terms(blocks)
        self.process_paragraphs(doc.paragraphs)
        self.process_tables(doc.tables)
        self.process_headers(doc)
        self.process_images(doc)
        output = path.parent / f"Anonymised_{path.name}"
        doc.save(output)

    # -------------------------
    # Process all DOCX in folder with multithreading
    # -------------------------
    def process_folder(self, folder):
        folder = Path(folder)
        if not folder.exists():
            print("Invalid folder path provided")
            return
        files = [f for f in folder.glob("*.docx") if not f.name.startswith("Anonymised_")]
        if not files:
            print("No DOCX files found to process")
            return

        # Determine thread pool size (75% of CPU cores)
        max_workers = max(1, int(multiprocessing.cpu_count() * 0.75))

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(self.anonymise_file, f): f for f in files}
            for f in tqdm(as_completed(futures), total=len(futures), desc="Anonymising reports"):
                pass

        print("\nAnonymisation complete")


def main():
    if len(sys.argv) < 2:
        print(
            "Usage:\n"
            "python pentest_report_anonymiser.py <folder>\n"
            "python pentest_report_anonymiser.py <folder> \"ClientName|ProjectName\""
        )
        return
    folder = sys.argv[1]
    custom_terms = None
    if len(sys.argv) >= 3:
        custom_terms = sys.argv[2]
    anonymiser = PentestAnonymiser(custom_terms)
    anonymiser.process_folder(folder)


if __name__ == "__main__":
    main()
